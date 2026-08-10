#!/usr/bin/env python3
"""
Build the capstone report .docx from the Markdown chapters in docs/report/.

    pip install python-docx
    python docs/build_report.py                 # -> docs/Capstone-Report.docx
    python docs/build_report.py -o out.docx

Applies the formatting the submission template mandates: Times New Roman,
12 pt body and 14 pt headings, 1.5 line spacing, one-inch margins, and page
numbers bottom-centre. The table of contents and the figure and table lists
are emitted as Word fields — open the result and press Ctrl+A then F9 to
populate them.

Cover-page details come from report-meta.json; anything left blank there is
rendered as an obvious underscore run so it cannot be missed in the PDF.

Only the Markdown this report actually uses is supported: ATX headings, fenced
code, pipe tables, bullet and numbered lists, blockquotes, horizontal rules,
and inline bold / italic / code / links. Adding syntax means extending
parse_blocks and render_block, not reaching for a Markdown library — keeping
python-docx as the single dependency is the point.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit(
        "python-docx is required.\n\n    pip install python-docx\n"
    )

REPORT_DIR = Path(__file__).parent / "report"
DEFAULT_OUT = Path(__file__).parent / "Capstone-Report.docx"

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BODY_PT = 12
HEADING_PT = 14
LINE_SPACING = 1.5

# Chapters are concatenated in filename order, which is why they are numbered.
CHAPTER_GLOB = "[0-9]*.md"


# ── metadata ────────────────────────────────────────────────────────────────

BLANK = "__________________"


def load_meta() -> dict:
    path = REPORT_DIR / "report-meta.json"
    if not path.exists():
        return {}
    meta = json.loads(path.read_text(encoding="utf-8"))
    meta.pop("_comment", None)
    return meta


def build_substitutions(meta: dict) -> dict:
    """Flatten report-meta.json into {{placeholder}} -> replacement text."""
    subs = {}
    for key, value in meta.items():
        if isinstance(value, (str, int, float)):
            subs[key] = str(value).strip() or BLANK

    students = meta.get("students") or []

    # Cover page: one line per student, with roll number where supplied.
    lines = []
    for s in students:
        name = (s.get("name") or "").strip() or BLANK
        roll = (s.get("roll") or "").strip() or BLANK
        lines.append(f"**{name}** — {roll}")
    subs["students_block"] = "\n\n".join(lines) if lines else BLANK

    # Declaration: a signature line per student.
    sig = []
    for s in students:
        name = (s.get("name") or "").strip() or BLANK
        sig.append(f"{name}    ________________________")
    subs["students_signature_block"] = "\n\n".join(sig) if sig else BLANK

    return subs


PLACEHOLDER_RE = re.compile(r"\{\{([a-z_]+)\}\}")


def substitute(text: str, subs: dict) -> str:
    def replace(match):
        key = match.group(1)
        if key in {"PAGEBREAK", "TOC", "LOF", "LOT"}:
            return match.group(0)  # structural, handled by the parser
        return subs.get(key, BLANK)

    return PLACEHOLDER_RE.sub(replace, text)


# ── markdown parsing ────────────────────────────────────────────────────────

FIGURE_RE = re.compile(r"^>\s*\*\*(Figure\s+[\w.]+)\s*[—-]\s*(.+?)\*\*", re.I)
TABLE_CAPTION_RE = re.compile(r"^\*\*(Table\s+[\w.]+)\s*[—-]\s*(.+?)\*\*\s*$", re.I)


def parse_blocks(text: str) -> list[dict]:
    """Turn Markdown into a flat list of block tokens."""
    blocks: list[dict] = []
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # HTML comments are editorial notes, not content.
        if stripped.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # Structural directives.
        if stripped in ("{{PAGEBREAK}}", "{{TOC}}", "{{LOF}}", "{{LOT}}"):
            blocks.append({"type": stripped.strip("{}").lower()})
            i += 1
            continue

        # Fenced code.
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1
            body = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                body.append(lines[i])
                i += 1
            i += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(body)})
            continue

        # Horizontal rule.
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            blocks.append({"type": "rule"})
            i += 1
            continue

        # Heading.
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading.group(1)),
                    "text": heading.group(2).strip(),
                }
            )
            i += 1
            continue

        # Pipe table: a header row followed by a separator row.
        if stripped.startswith("|") and i + 1 < len(lines) and _is_separator(lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not _is_separator(lines[i]):
                    rows.append(_split_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "rows": rows})
            continue

        # Blockquote — used in this report for figure placeholders.
        if stripped.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            joined = " ".join(q for q in quote if q)
            figure = FIGURE_RE.match("> **" + joined + "**" if not joined.startswith("**") else "> " + joined)
            if figure:
                blocks.append(
                    {"type": "caption", "kind": "Figure",
                     "label": figure.group(1), "title": figure.group(2).rstrip(".")}
                )
            else:
                blocks.append({"type": "quote", "text": joined})
            continue

        # Lists.
        bullet = re.match(r"^([-*+])\s+(.*)$", stripped)
        numbered = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if bullet or numbered:
            ordered = numbered is not None
            items = []
            while i < len(lines):
                current = lines[i].strip()
                match = (
                    re.match(r"^(\d+)[.)]\s+(.*)$", current)
                    if ordered
                    else re.match(r"^([-*+])\s+(.*)$", current)
                )
                if not match:
                    # A blank line inside a list is a separator, not a terminator.
                    if not current and i + 1 < len(lines):
                        nxt = lines[i + 1].strip()
                        still = (
                            re.match(r"^\d+[.)]\s+", nxt) if ordered
                            else re.match(r"^[-*+]\s+", nxt)
                        )
                        if still:
                            i += 1
                            continue
                    break
                items.append(match.group(2).strip())
                i += 1
            blocks.append({"type": "list", "ordered": ordered, "items": items})
            continue

        # Standalone table caption.
        caption = TABLE_CAPTION_RE.match(stripped)
        if caption:
            blocks.append(
                {"type": "caption", "kind": "Table",
                 "label": caption.group(1), "title": caption.group(2).rstrip(".")}
            )
            i += 1
            continue

        # Paragraph: consume until a blank line or a new block starts.
        para = []
        while i < len(lines):
            current = lines[i].strip()
            if not current or _starts_block(current):
                break
            para.append(current)
            i += 1
        if para:
            blocks.append({"type": "paragraph", "text": " ".join(para)})

    return blocks


def _is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s|:-]+\|?", line.strip()))


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _starts_block(line: str) -> bool:
    return bool(
        line.startswith(("#", ">", "|", "```", "{{"))
        or re.match(r"^[-*+]\s+", line)
        or re.match(r"^\d+[.)]\s+", line)
        or re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line)
    )


# ── inline formatting ───────────────────────────────────────────────────────

INLINE_RE = re.compile(
    r"(\*\*.+?\*\*)"      # bold
    r"|(`[^`]+`)"         # code
    r"|(\*[^*]+\*)"       # italic
    r"|(\[[^\]]+\]\([^)]+\))"  # link
)


def add_inline(paragraph, text: str) -> None:
    """Render inline markdown into runs on an existing paragraph."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(BODY_PT - 1.5)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("[") and "](" in part:
            label, _, url = part[1:-1].partition("](")
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0x1F, 0x3D, 0x99)
            run.underline = True
            del url  # the printed report is the deliverable; the label carries it
        else:
            paragraph.add_run(part)


# ── OOXML helpers ───────────────────────────────────────────────────────────

def _field(paragraph, instruction: str, placeholder: str) -> None:
    """Insert a Word field. Populated when the reader updates fields (F9)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def add_page_numbers(section) -> None:
    """Bottom-centre page number, as the template requires."""
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    _field(paragraph, "PAGE", "1")
    for run in paragraph.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_PT - 2)


def add_seq_caption(doc, kind: str, label: str, title: str):
    """
    A caption carrying a SEQ field, which is what lets Word build the figure
    and table lists. Uses the built-in Caption style so the TOC \\c switch
    picks it up.
    """
    try:
        paragraph = doc.add_paragraph(style="Caption")
    except KeyError:
        paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"{kind} ")
    _field(paragraph, f"SEQ {kind} \\* ARABIC", label.split()[-1])
    paragraph.add_run(f" — {title}")
    for run in paragraph.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_PT - 1)
        run.italic = True
    return paragraph


# ── document setup ──────────────────────────────────────────────────────────

def configure(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    # East-Asian font must be set too, or Word substitutes for some glyphs.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing = LINE_SPACING
    normal.paragraph_format.space_after = Pt(6)

    for level in range(1, 5):
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        style.font.size = Pt(HEADING_PT if level <= 2 else BODY_PT)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = LINE_SPACING
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_page_numbers(section)


def add_code_block(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.0
    fmt.left_indent = Inches(0.3)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(9.5)

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "F4F4F4")
    paragraph._p.get_or_add_pPr().append(shading)


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for column in range(columns):
            content = row[column] if column < len(row) else ""
            paragraph = cells[column].paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline(paragraph, content)
            for run in paragraph.runs:
                run.font.size = Pt(BODY_PT - 2)
                if index == 0:
                    run.bold = True


def render_block(doc, block: dict) -> None:
    kind = block["type"]

    if kind == "pagebreak":
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    elif kind == "toc":
        doc.add_heading("Table of Contents", level=1)
        paragraph = doc.add_paragraph()
        _field(paragraph, r'TOC \o "1-3" \h \z \u',
               "Select all (Ctrl+A) and press F9 to populate.")

    elif kind == "lof":
        doc.add_heading("List of Figures", level=1)
        paragraph = doc.add_paragraph()
        _field(paragraph, r'TOC \h \z \c "Figure"',
               "Select all (Ctrl+A) and press F9 to populate.")

    elif kind == "lot":
        doc.add_heading("List of Tables", level=1)
        paragraph = doc.add_paragraph()
        _field(paragraph, r'TOC \h \z \c "Table"',
               "Select all (Ctrl+A) and press F9 to populate.")

    elif kind == "heading":
        doc.add_heading(block["text"], level=min(block["level"], 4))

    elif kind == "caption":
        add_seq_caption(doc, block["kind"], block["label"], block["title"])

    elif kind == "code":
        add_code_block(doc, block["text"])

    elif kind == "table":
        add_table(doc, block["rows"])

    elif kind == "rule":
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(2)
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "BBBBBB")
        border.append(bottom)
        paragraph._p.get_or_add_pPr().append(border)

    elif kind == "quote":
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.4)
        add_inline(paragraph, block["text"])
        for run in paragraph.runs:
            run.italic = True

    elif kind == "list":
        style = "List Number" if block["ordered"] else "List Bullet"
        for item in block["items"]:
            try:
                paragraph = doc.add_paragraph(style=style)
            except KeyError:
                paragraph = doc.add_paragraph()
                item = f"• {item}"
            paragraph.paragraph_format.line_spacing = LINE_SPACING
            # Markdown task-list syntax has no Word equivalent; use symbols.
            item = re.sub(r"^\[ \]\s*", "☐ ", item)
            item = re.sub(r"^\[[xX]\]\s*", "☑ ", item)
            add_inline(paragraph, item)

    elif kind == "paragraph":
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(paragraph, block["text"])


# ── entry point ─────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("-o", "--output", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()

    chapters = sorted(REPORT_DIR.glob(CHAPTER_GLOB))
    if not chapters:
        sys.exit(f"No chapters found in {REPORT_DIR}")

    meta = load_meta()
    subs = build_substitutions(meta)

    missing = sorted(k for k, v in subs.items() if v == BLANK)
    if missing:
        print("Placeholders still unfilled in report-meta.json:")
        for key in missing:
            print(f"  - {key}")
        print()

    doc = Document()
    configure(doc)

    total_blocks = 0
    for path in chapters:
        text = substitute(path.read_text(encoding="utf-8"), subs)
        blocks = parse_blocks(text)
        total_blocks += len(blocks)
        for block in blocks:
            render_block(doc, block)
        print(f"  {path.name:28} {len(blocks):4} blocks")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)

    print(f"\nWrote {args.output} ({total_blocks} blocks from {len(chapters)} chapters)")
    print("\nIn Word: Ctrl+A then F9 to populate the contents and figure lists,")
    print("then File -> Export -> Create PDF/XPS to produce the submission file.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
